import os, uuid, threading, re
from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
from deep_translator import GoogleTranslator
from docx import Document

app = Flask(__name__)
CORS(app)

UPLOAD_DIR = "uploads"
OUTPUT_DIR = "outputs"
os.makedirs(UPLOAD_DIR, exist_ok=True)
os.makedirs(OUTPUT_DIR, exist_ok=True)

translator = GoogleTranslator(source="en", target="hi")

jobs = {}

# -------- Advanced Cleaning --------
def clean_text(text):
    # Remove OCR garbage (random mixed characters like hetJe&), but preserve punctuation
    # This regex looks for words containing special chars in the middle
    text = re.sub(r"\b\w*[&@#]+\w*\b", "", text)
    # Remove extra whitespace
    text = re.sub(r"\s{2,}", " ", text)
    return text.strip()

# -------- Smart Translation Logic --------
def smart_translate(text):
    if not text:
        return ""

    # REGEX EXPLANATION:
    # We look for a pattern at the VERY END ($) of the string.
    # It looks for an opening parenthesis '(', followed by any text, 
    # then the word 'Exam' or 'Shift' (to be safe), and a closing parenthesis ')'.
    # This captures: (MP Police Constable Exam 23/07/2016, Shift-I)
    
    pattern = r"(\s*\(.*?Exam.*?\))$" 
    
    match = re.search(pattern, text, re.IGNORECASE)

    if match:
        # 1. Extract the English Tag (Don't translate this)
        english_tag = match.group(1) 
        
        # 2. Extract the Content part (Translate this)
        # We take everything from start up to where the match began
        content_part = text[:match.start()]
        
        # 3. Translate content
        if content_part.strip():
            try:
                # Translate only the question text
                translated_content = translator.translate(content_part.strip())
            except Exception as e:
                print(f"Translation Error: {e}")
                translated_content = content_part
        else:
            translated_content = ""

        # 4. Merge them back together
        # Result: "Hindi Text" + " (English Tag)"
        return f"{translated_content} {english_tag.strip()}"

    else:
        # If no exam tag is found, translate the whole line normally
        try:
            return translator.translate(text)
        except:
            return text

# -------- Background job --------
def process_doc(job_id, input_path, output_path):
    doc = Document(input_path)
    out = Document()

    total = len(doc.paragraphs)
    jobs[job_id]["total"] = total

    for i, para in enumerate(doc.paragraphs, start=1):
        raw_text = para.text.strip()
        
        # 1. Clean garbage first
        cleaned = clean_text(raw_text)

        jobs[job_id]["current"] = i
        jobs[job_id]["line"] = (cleaned[:50] + '...') if len(cleaned) > 50 else cleaned 
        
        print(f"[JOB {job_id}] Processing: {cleaned}")

        if cleaned:
            # 2. Use the split-translate-merge logic
            final_text = smart_translate(cleaned)
            out.add_paragraph(final_text)
        else:
            out.add_paragraph("")

    out.save(output_path)
    jobs[job_id]["done"] = True
    jobs[job_id]["file"] = output_path
    print(f"[JOB {job_id}] DONE")

# -------- Routes --------
@app.route("/upload", methods=["POST"])
def upload():
    file = request.files.get("file")
    if not file or not file.filename.endswith(".docx"):
        return jsonify(error="Only DOCX allowed"), 400

    job_id = str(uuid.uuid4())
    input_path = f"{UPLOAD_DIR}/{job_id}.docx"
    output_path = f"{OUTPUT_DIR}/{job_id}_hi.docx"

    file.save(input_path)

    jobs[job_id] = {
        "current": 0,
        "total": 0,
        "line": "",
        "done": False,
        "file": None
    }

    threading.Thread(
        target=process_doc,
        args=(job_id, input_path, output_path),
        daemon=True
    ).start()

    return jsonify(job_id=job_id)

@app.route("/status/<job_id>")
def status(job_id):
    return jsonify(jobs.get(job_id, {}))

@app.route("/download/<job_id>")
def download(job_id):
    job = jobs.get(job_id)
    if job and job["done"]:
        return send_file(job["file"], as_attachment=True)
    return jsonify(error="Not ready"), 400

if __name__ == "__main__":
    port = int(os.environ.get("PORT", 10000))
    app.run(host="0.0.0.0", port=port)
